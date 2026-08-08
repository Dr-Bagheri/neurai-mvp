"""Minutes & export (D5): meeting templates incl. the official Iranian
صورتجلسه, plus SRT subtitles. Word export needs python-docx (optional extra);
SRT/TXT are dependency-free.

The صورتجلسه body is produced by the LLM under the same consent constraints
as everything else, then register-converted (گفتاری→نوشتاری) — dates are
Jalali."""
from __future__ import annotations

from datetime import date, datetime
from pathlib import Path

from neurai.config import get_config
from neurai.db import get_db
from neurai.fa import fa_normalize, jalali_str
from neurai.harness import get_harness
from neurai.harness.harness import Constraints
from neurai.rag.ingest import transcript_text


def _ts(ms: int) -> str:
    h, rem = divmod(ms, 3_600_000)
    m, rem = divmod(rem, 60_000)
    s, msec = divmod(rem, 1_000)
    return f"{h:02d}:{m:02d}:{s:02d},{msec:03d}"


def build_srt(meeting_id: int) -> str:
    db = get_db()
    for pass_name in ("quality", "live"):
        rows = db.query(
            "SELECT start_ms, end_ms, speaker_label, text FROM transcript_segments "
            "WHERE meeting_id=? AND pass=? ORDER BY start_ms",
            (meeting_id, pass_name),
        )
        if rows:
            break
    names = {
        r["label"]: r["display_name"]
        for r in db.query("SELECT label, display_name FROM speaker_names WHERE meeting_id=?", (meeting_id,))
    }
    blocks = []
    for i, r in enumerate(rows, start=1):
        speaker = names.get(r["speaker_label"], r["speaker_label"])
        prefix = f"{speaker}: " if speaker else ""
        blocks.append(f"{i}\n{_ts(r['start_ms'])} --> {_ts(r['end_ms'])}\n{prefix}{r['text']}\n")
    return "\n".join(blocks)


_TEMPLATE_PROMPTS = {
    "soorat_jalase": (
        "از رونوشت جلسه، متن بخش‌های صورت‌جلسه رسمی را به فارسی نوشتاری تهیه کن. "
        "خروجی را دقیقاً با این ساختار بده:\n"
        "## دستور جلسه\n…\n## خلاصه مذاکرات\n…\n## مصوبات\n"
        "(هر مصوبه در یک سطر: شرح، مسئول پیگیری، مهلت)\n"
        "زبان باید کاملاً رسمی و اداری باشد."
    ),
    "standup": (
        "از رونوشت این جلسه استندآپ، برای هر نفر سه بخش بنویس: انجام‌شده، در دست انجام، موانع. "
        "کوتاه و فهرست‌وار، فارسی رسمی."
    ),
    "plain": "از رونوشت جلسه صورت‌جلسه‌ای ساده و ساخت‌یافته به فارسی رسمی بنویس.",
}


async def build_minutes_body(meeting_id: int, allow_cloud: bool, template: str) -> tuple[str, str]:
    """Returns (body_markdown, source)."""
    text = transcript_text(meeting_id)
    if not text:
        raise RuntimeError("رونوشتی برای این جلسه موجود نیست")
    prompt = _TEMPLATE_PROMPTS.get(template, _TEMPLATE_PROMPTS["plain"])
    result = await get_harness().summarize_long(
        "minutes", prompt, text, Constraints(allow_cloud=allow_cloud),
    )
    return fa_normalize(result.text), result.source


def _header_lines(meeting) -> list[str]:
    started = meeting["started_at"] or meeting["created_at"]
    try:
        day = datetime.fromisoformat(started).date()
    except (ValueError, TypeError):
        day = date.today()
    return [
        "صورت‌جلسه",
        f"موضوع جلسه: {meeting['title']}",
        f"تاریخ: {jalali_str(day)}",
        "حاضرین: ...........................",
        "غایبین: ...........................",
    ]


def _signature_lines() -> list[str]:
    return ["", "امضاء حاضرین:", "", ".................    .................    ................."]


async def export_meeting(meeting_id: int, ctx, fmt: str = "docx", template: str = "soorat_jalase") -> Path:
    cfg = get_config()
    db = get_db()
    meeting = db.query_one("SELECT * FROM meetings WHERE id=?", (meeting_id,))
    if meeting is None:
        raise RuntimeError("جلسه پیدا نشد")

    out_dir = cfg.exports_dir / str(meeting_id)
    out_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")

    if fmt == "srt":
        path = out_dir / f"meeting-{meeting_id}-{stamp}.srt"
        path.write_text(build_srt(meeting_id), encoding="utf-8")
        return path

    # D3: cloud needs the caller's consent AND the meeting's own opt-in
    body, _source = await build_minutes_body(
        meeting_id, ctx.allow_cloud and bool(meeting["allow_cloud"]), template,
    )
    header = _header_lines(meeting)
    footer = _signature_lines()

    if fmt == "txt":
        path = out_dir / f"minutes-{meeting_id}-{stamp}.txt"
        path.write_text("\n".join(header) + "\n\n" + body + "\n" + "\n".join(footer), encoding="utf-8")
        return path

    if fmt == "docx":
        try:
            import docx  # optional extra: pip install neurai-server[export]
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as e:
            raise RuntimeError("خروجی Word نیاز به python-docx دارد (pip install python-docx)") from e

        doc = docx.Document()
        # RTL: set paragraph alignment right; full RTL styling is a webui/print concern
        title = doc.add_heading(header[0], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for line in header[1:]:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for line in body.splitlines():
            if line.startswith("## "):
                h = doc.add_heading(line[3:], level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif line.strip():
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for line in footer:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        path = out_dir / f"minutes-{meeting_id}-{stamp}.docx"
        doc.save(str(path))
        return path

    raise RuntimeError(f"unsupported format: {fmt}")
